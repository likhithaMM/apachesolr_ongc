import os
import pytesseract
from pdf2image import convert_from_path
from docx import Document

# Path to directory containing PDF and DOC files
directory = 'content'

# Output directory to save extracted text files
output_directory = 'extracted_text'

# Create the output directory if it doesn't exist
if not os.path.exists(output_directory):
    os.makedirs(output_directory)

# Loop through files in the directory
for filename in os.listdir(directory):
    if filename.endswith('.pdf'):
        print("Processing PDF file:", filename)
        # Convert PDF to images
        pages = convert_from_path(os.path.join(directory, filename))
        # OCR extraction from images
        extracted_text = ''
        for page_num, page in enumerate(pages):
            text = pytesseract.image_to_string(page)
            extracted_text += text + '\n\n'  # Append extracted text from each page
        # Save extracted text to a text file
        output_filename = os.path.splitext(filename)[0] + '_extracted.txt'
        with open(os.path.join(output_directory, output_filename), 'w', encoding='utf-8') as f:
            f.write(extracted_text)
        print("Extracted text saved to:", output_filename)
    elif filename.endswith('.doc'):
        print("Processing DOC file:", filename)
        doc = Document(os.path.join(directory, filename))
        extracted_text = ''
        for paragraph in doc.paragraphs:
            extracted_text += paragraph.text + '\n\n'  # Append extracted text from each paragraph
        # Save extracted text to a text file
        output_filename = os.path.splitext(filename)[0] + '_extracted.txt'
        with open(os.path.join(output_directory, output_filename), 'w', encoding='utf-8') as f:
            f.write(extracted_text)
        print("Extracted text saved to:", output_filename)





# import os
# import pytesseract
# from pdf2image import convert_from_path
# from docx import Document

# # Path to directory containing PDF and DOC files
# directory = 'content'

# # Loop through files in the directory
# for filename in os.listdir(directory):
#     if filename.endswith('.pdf'):
#         print("Processing PDF file:", filename)
#         # Convert PDF to images
#         pages = convert_from_path(os.path.join(directory, filename))
#         # OCR extraction from images
#         for page_num, page in enumerate(pages):
#             text = pytesseract.image_to_string(page)
#             print("Extracted text from page", page_num + 1, ":", text)
#     elif filename.endswith('.doc'):
#         print("Processing DOC file:", filename)
#         doc = Document(os.path.join(directory, filename))
#         text = ""
#         for paragraph in doc.paragraphs:
#             text += paragraph.text
#         print("Extracted text:", text)
